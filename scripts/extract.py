#!/usr/bin/env python3
"""
题目提取脚本 — 支持PDF(文字型/扫描型)、Word、图片
自动检测PDF类型，文字型用pdfplumber，扫描型用PaddleOCR+视觉AI
"""

import argparse
import os
import sys
import json
import re


def extract_text_pdf(pdf_path, output_dir=None, verbose=False):
    """文字型PDF：用pdfplumber直接提取文字"""
    import pdfplumber

    results = []
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            results.append({
                "page": i + 1,
                "text": text.strip(),
                "type": "text"
            })
            if verbose:
                print(f"  Page {i+1}: {len(text.strip())} chars extracted")

    return results


def extract_scanned_pdf(pdf_path, output_dir=None, verbose=False):
    """扫描型PDF：用PaddleOCR提取文字，同时保存图片供视觉AI读取公式/电路图"""
    import fitz  # PyMuPDF

    # 尝试导入PaddleOCR
    try:
        from paddleocr import PaddleOCR
        ocr = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)
        has_ocr = True
        if verbose:
            print("  PaddleOCR loaded successfully")
    except Exception as e:
        has_ocr = False
        if verbose:
            print(f"  PaddleOCR not available: {e}")
            print("  Falling back to image-only mode for visual AI")

    doc = fitz.open(pdf_path)
    results = []

    # 确保图片目录存在
    if output_dir:
        img_dir = os.path.join(output_dir, "图片")
        os.makedirs(img_dir, exist_ok=True)

    for i, page in enumerate(doc):
        page_num = i + 1

        # 保存页面图片
        img_path = None
        if output_dir:
            pix = page.get_pixmap(dpi=300)
            img_path = os.path.join(img_dir, f"page_{page_num:03d}.png")
            pix.save(img_path)

        # OCR提取文字
        ocr_text = ""
        if has_ocr and img_path:
            try:
                result = ocr.ocr(img_path, cls=True)
                if result and result[0]:
                    lines = [line[1][0] for line in result[0]]
                    ocr_text = "\n".join(lines)
            except Exception as e:
                if verbose:
                    print(f"  Page {page_num} OCR failed: {e}")

        results.append({
            "page": page_num,
            "text": ocr_text.strip(),
            "image": img_path,
            "type": "scanned"
        })

        if verbose:
            status = f"{len(ocr_text.strip())} chars" if ocr_text else "image only"
            print(f"  Page {page_num}: {status}")

    doc.close()
    return results


def detect_pdf_type(pdf_path, verbose=False):
    """检测PDF是文字型还是扫描型"""
    import pdfplumber

    with pdfplumber.open(pdf_path) as pdf:
        # 检查前3页
        sample_pages = min(3, len(pdf.pages))
        total_chars = 0

        for i in range(sample_pages):
            text = pdf.pages[i].extract_text() or ""
            total_chars += len(text.strip())

        avg_chars = total_chars / sample_pages if sample_pages > 0 else 0

        if verbose:
            print(f"  Sampled {sample_pages} pages, avg {avg_chars:.0f} chars/page")

        # 少于50字/页认为是扫描型
        is_text = avg_chars >= 50
        return "text" if is_text else "scanned"


def extract_pdf(pdf_path, output_dir=None, verbose=False):
    """提取PDF题目，自动检测类型"""
    pdf_type = detect_pdf_type(pdf_path, verbose)
    if verbose:
        print(f"  Detected type: {pdf_type}")

    if pdf_type == "text":
        return extract_text_pdf(pdf_path, output_dir, verbose)
    else:
        return extract_scanned_pdf(pdf_path, output_dir, verbose)


def extract_docx(docx_path, verbose=False):
    """Word文档：用python-docx提取文字"""
    from docx import Document

    doc = Document(docx_path)
    full_text = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)

    # 也提取表格中的文字
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_text.append(row_text)

    result = "\n".join(full_text)

    if verbose:
        print(f"  Extracted {len(result)} chars from docx")

    return [{
        "page": 1,
        "text": result,
        "type": "docx"
    }]


def extract_image(img_path, verbose=False):
    """图片：返回路径供视觉AI处理"""
    if verbose:
        print(f"  Image path: {img_path}")

    return [{
        "page": 1,
        "text": "",  # 视觉AI需要单独处理
        "image": os.path.abspath(img_path),
        "type": "image"
    }]


def extract(file_path, output_dir=None, verbose=False):
    """统一入口：根据文件类型自动选择提取方式"""
    file_path = os.path.abspath(file_path)

    if not os.path.exists(file_path):
        print(f"Error: File not found: {file_path}")
        sys.exit(1)

    ext = os.path.splitext(file_path)[1].lower()

    if verbose:
        print(f"Extracting: {file_path}")
        print(f"File type: {ext}")

    if ext == ".pdf":
        return extract_pdf(file_path, output_dir, verbose)
    elif ext in (".docx", ".doc"):
        return extract_docx(file_path, verbose)
    elif ext in (".png", ".jpg", ".jpeg", ".bmp", ".tiff"):
        return extract_image(file_path, verbose)
    else:
        print(f"Error: Unsupported file type: {ext}")
        print("Supported: .pdf, .docx, .doc, .png, .jpg, .jpeg, .bmp, .tiff")
        sys.exit(1)


def save_results(results, output_path):
    """保存提取结果到txt文件"""
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    with open(output_path, "w", encoding="utf-8") as f:
        for r in results:
            f.write(f"=== Page {r['page']} [{r['type']}] ===\n")
            if r.get("text"):
                f.write(r["text"] + "\n")
            if r.get("image"):
                f.write(f"[Image: {r['image']}]\n")
            f.write("\n")

    print(f"Saved to: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="提取PDF/Word/图片中的题目内容")
    parser.add_argument("file", help="输入文件路径（PDF/Word/图片）")
    parser.add_argument("-o", "--output", help="输出txt文件路径")
    parser.add_argument("-d", "--dir", help="项目目录（用于保存图片）")
    parser.add_argument("--type", choices=["text", "scanned", "auto"],
                        default="auto", help="PDF类型（默认自动检测）")
    parser.add_argument("--json", action="store_true", help="输出JSON格式")
    parser.add_argument("-v", "--verbose", action="store_true", help="详细输出")

    args = parser.parse_args()

    # 提取
    if args.type == "auto":
        results = extract(args.file, args.dir, args.verbose)
    elif args.type == "text":
        results = extract_text_pdf(args.file, args.dir, args.verbose)
    elif args.type == "scanned":
        results = extract_scanned_pdf(args.file, args.dir, args.verbose)

    # 输出
    if args.json:
        print(json.dumps(results, ensure_ascii=False, indent=2))
    elif args.output:
        save_results(results, args.output)
    else:
        # 默认输出到同目录
        base = os.path.splitext(os.path.basename(args.file))[0]
        out_dir = args.dir or os.path.dirname(args.file)
        out_path = os.path.join(out_dir, "OCR", f"{base}.txt")
        save_results(results, out_path)


if __name__ == "__main__":
    main()
